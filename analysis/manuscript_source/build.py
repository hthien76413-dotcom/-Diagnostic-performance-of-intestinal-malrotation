import docx, json, re, os
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
AN='/tmp/claude-0/-home-user--Diagnostic-performance-of-intestinal-malrotation/4998e6b2-e14c-57b0-80b5-41a1fb987d43/scratchpad/an/'
OUT='/home/user/-Diagnostic-performance-of-intestinal-malrotation/'
T=json.load(open(AN+'tables123.json')); T.update(json.load(open(AN+'tables456.json')))
TITLES={
 'T1':('Table 1','Characteristics of the 465 children with surgically confirmed intestinal malrotation, overall and according to which preoperative index test they received',
   'Groups overlap, because a child could receive more than one index test; columns therefore do not sum to the cohort total. Age at operation was calculated as age at admission for the operative encounter plus the interval from admission to operation. Presenting features were extracted from admission records by text search and are documentation rates, not verified prevalences. The rightmost column shows the 55 children who received none of the three index tests; all of them nonetheless had other preoperative imaging (Fig. 1). IQR interquartile range, UGI upper gastrointestinal.'),
 'T2':('Table 2','Report-level detection of intestinal malrotation among surgically confirmed children, with the certainty of the wording used in positive conclusions',
   'Wilson 95% confidence intervals. Denominators differ between modalities and are drawn from overlapping but non-identical, indication-selected groups of children; the rates are not directly comparable between modalities and are not sensitivities. Certainty tiers were assigned from the conclusion text: definite (unqualified statement), probable ("most likely", "first consideration"), possible ("suspected", "cannot be excluded", "?"). The final column repeats the detection rate after reclassifying all possible-tier conclusions as negative. Percentages for certainty tiers are of positive reports; detection rates are of all index reports of that modality. Contrast-enhanced and unenhanced CT were performed for different indications and in children of different ages, so their comparison is confounded and is presented as an exploratory subgroup only.'),
 'T3':('Table 3','Documented content of the 119 routine gastrointestinal ultrasound index reports, and report-level detection conditional on that content',
   'Content was coded from the findings and conclusion text of each index report by pre-specified patterns (Online Resource 1). This is an audit of what was recorded, and is a lower bound on what was performed: an element assessed but not documented cannot be distinguished here from one never assessed. The whirlpool sign appears twice because the adjudicated study variable and the text-pattern variable differ slightly (59 vs 57 reports); both are shown for transparency. Examination-type categories are not mutually exclusive.'),
 'T4':('Table 4','Between-modality comparison of report-level detection (generalised estimating equation)',
   'GEE logistic model with exchangeable working correlation and patient-level clustering (cluster-robust standard errors); 410 children, 740 preoperative index examinations. An odds ratio (OR) below 1 indicates lower report-level detection than the UGI series. These estimates describe indication-driven detection under routine test selection and are not estimates of comparative diagnostic accuracy; the modality coefficients absorb the indication for the test, its position in the diagnostic pathway and the content of the examination. The pre-specified modality-by-volvulus interaction could not be estimated across all three modalities because no ultrasound examination was positive among the six children without volvulus (complete separation, model non-convergence); for the estimable UGI-versus-CT comparison the interaction OR was 0.90 (95% CI 0.33-2.43, p=0.84).'),
 'T5':('Table 5','Report-level detection in the selected subgroup of children who underwent all three examinations, overall and restricted to examinations performed close together in time',
   'This subgroup was assembled by diagnostic uncertainty, not by sampling: 96.6% had midgut volvulus, 83.1% were neonates and 62.7% were operated on in 2019-2026, compared with 86.6%, 65.5% and 31.3% of the other imaged children. It answers the question of which examination named the diagnosis most often among children investigated intensively enough to receive all three, and is not a population-level comparison of test accuracy. The interval is that between the first and last of the three examinations (median 0.9 days, IQR 0.6-1.6). Discordant pairs are shown as first-positive/second-positive.'),
 'T6':('Table 6','Change in report-level detection between eras, and the examination-content variables that account for it',
   'Modality-specific logistic models. For ultrasound the content variable is whether the examination was recorded as an abdominal great-vessel study rather than as gastrointestinal ultrasound; for CT it is intravenous contrast enhancement. Attenuation of the era odds ratio after the content variable is added indicates that the temporal change operated through what the examination was rather than through calendar time. The content variables were not randomly assigned and are themselves confounded by indication; these models are explanatory rather than causal.'),
}
FIGS={'FIG1':('Fig. 1',OUT+'Figure1_flow.png','Study flow diagram showing assembly of the surgically confirmed malrotation cohort, availability of each preoperative index test, and the imaging actually received by the 55 children who had none of the three index examinations',6.4),
 'FIG2':('Fig. 2',OUT+'Figure2_paired.png','Report-level detection in the selected subgroup of 59 children who underwent all three examinations preoperatively, shown for the whole subgroup and for the subsets in which all three examinations fell within 48 h and within 24 h. Error bars are Wilson 95% confidence intervals. This subgroup was assembled by diagnostic uncertainty and is not a population-level comparison of test accuracy',6.4),
 'FIG3':('Fig. 3',OUT+'Figure3_detection.png','Report-level detection of malrotation by each modality among children with a preoperative index examination (Wilson 95% confidence intervals). Each estimate has its own denominator, drawn from a different, indication-selected group of children; the rates are not directly comparable between modalities and do not constitute sensitivity',6.4),
 'FIG4':('Fig. 4',OUT+'Figure4_us_audit.png','Content of the 119 routine ultrasound index reports (a) and report-level detection conditional on that content (b). No report documented the third portion of the duodenum or the duodenojejunal junction; detection was 100% when a whirlpool sign was recorded and 10% when it was not',6.6)}

doc=docx.Document()
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(11)
for s in doc.sections:
    s.left_margin=s.right_margin=Inches(1.0)
def para(text,style=None,italic=False,size=None,align=None,space_after=8):
    p=doc.add_paragraph(style=style)
    # bold **...**
    parts=re.split(r'(\*\*[^*]+\*\*)',text)
    for pt in parts:
        if not pt: continue
        if pt.startswith('**') and pt.endswith('**'):
            r=p.add_run(pt[2:-2]); r.bold=True
        else:
            r=p.add_run(pt)
        r.italic=italic
        if size: r.font.size=Pt(size)
    p.paragraph_format.space_after=Pt(space_after)
    if align: p.alignment=align
    return p
def add_table(key):
    tag,title,foot=TITLES[key]; data=T[key]
    p=para(f'{tag}. {title}'); p.runs[0].bold=True
    t=doc.add_table(rows=len(data),cols=len(data[0])); t.style='Light Grid Accent 1'
    for i,row in enumerate(data):
        for j,c in enumerate(row):
            cell=t.cell(i,j); cell.text=''
            r=cell.paragraphs[0].add_run(str(c)); r.font.size=Pt(8.5); r.font.name='Times New Roman'
            if i==0: r.bold=True
    para(foot,italic=True,size=8.5)
def add_fig(key):
    tag,path,cap,w=FIGS[key]
    if os.path.exists(path):
        doc.add_picture(path,width=Inches(w))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=para(f'{tag} {cap}',size=9.5)
    p.runs[0].bold=False

text=''
for f in ['part1.md','part2.md','part3.md','part4.md']:
    text+=open(f).read()+'\n'
lines=[l for l in text.split('\n')]
for ln in lines:
    ln=ln.strip()
    if not ln: continue
    if ln.startswith('#T '): 
        p=para(ln[3:],style='Title')
    elif ln.startswith('#H1 '):
        doc.add_heading(ln[4:],level=1)
    elif ln.startswith('#H2 '):
        doc.add_heading(ln[4:],level=2)
    elif ln.startswith('#N '):
        para(ln[3:])
    elif ln.startswith('#R '):
        para(ln[3:],size=10,space_after=3)
    elif ln.startswith('#TAB'):
        add_table('T'+ln[4:])
    elif ln.startswith('#FIG'):
        add_fig(ln[1:])
    else:
        para(ln)
# word count note after title block
doc.save(OUT+'诊断效能_英文稿_修回版_v3.docx')
print('saved')
