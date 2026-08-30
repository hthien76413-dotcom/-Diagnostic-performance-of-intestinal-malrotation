import docx, json, re
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
AN='/tmp/claude-0/-home-user--Diagnostic-performance-of-intestinal-malrotation/4998e6b2-e14c-57b0-80b5-41a1fb987d43/scratchpad/an/'
OUT='/home/user/-Diagnostic-performance-of-intestinal-malrotation/'
TB=json.load(open(AN+'or_tables.json')); TB.update(json.load(open(AN+'or_h.json')))
CAP={'H':('Table S1','Report-content audit patterns'),
     'S1':('Table S2','Distribution of algorithmic labels and certainty tiers, by modality'),
     'S2':('Table S3','Report-level detection stratified by midgut volvulus and by age category'),
     'S2b':('Table S4','Detection of a volvulus-specific sign among children with surgically confirmed midgut volvulus'),
     'S3':('Table S5','Documented content of the CT and upper gastrointestinal series index reports')}
def make(src,outfile):
    doc=docx.Document()
    st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(11)
    for s in doc.sections: s.left_margin=s.right_margin=Inches(1.0)
    def para(t,style=None,size=None,italic=False,space_after=8):
        p=doc.add_paragraph(style=style)
        for pt in re.split(r'(\*\*[^*]+\*\*)',t):
            if not pt: continue
            if pt.startswith('**') and pt.endswith('**'):
                r=p.add_run(pt[2:-2]); r.bold=True
            else: r=p.add_run(pt)
            r.italic=italic
            if size: r.font.size=Pt(size)
        p.paragraph_format.space_after=Pt(space_after); return p
    def table(key):
        tag,title=CAP[key]; data=TB[key]
        p=para(f'{tag}. {title}'); p.runs[0].bold=True
        t=doc.add_table(rows=len(data),cols=len(data[0])); t.style='Light Grid Accent 1'
        for i,row in enumerate(data):
            for j,c in enumerate(row):
                cell=t.cell(i,j); cell.text=''
                r=cell.paragraphs[0].add_run(str(c)); r.font.size=Pt(8.5); r.font.name='Times New Roman'
                if i==0: r.bold=True
        para('',space_after=4)
    for ln in open(src).read().split('\n'):
        ln=ln.strip()
        if not ln: continue
        if ln.startswith('#T '): para(ln[3:],style='Title')
        elif ln.startswith('#H1 '): doc.add_heading(ln[4:],level=1)
        elif ln.startswith('#H2 '): doc.add_heading(ln[4:],level=2)
        elif ln.startswith('#N '): para(ln[3:])
        elif ln=='#PB': doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        elif ln.startswith('#TAB'): table(ln[4:] if ln[4:] in TB else ln[4:].upper().replace('S2B','S2b'))
        else: para(ln)
    doc.save(OUT+outfile); print('saved',outfile)
make('or1.md','Online_Resource_1_NLP_and_report_audit.docx')
make('or23.md','Online_Resource_2_and_3_supplementary_tables.docx')
