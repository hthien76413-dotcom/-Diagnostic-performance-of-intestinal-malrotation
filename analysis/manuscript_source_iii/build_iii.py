import docx, json, re, os
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from threeline import make_three_line_table

def _blank_props(doc):
    """Strip python-docx's default authorship before saving.

    These files go out for double-blind review, so no document property may
    name an author, a reviser or the generator.
    """
    c = doc.core_properties
    for f in ('author', 'last_modified_by', 'title', 'subject', 'comments',
              'category', 'keywords', 'content_status', 'identifier',
              'language', 'version'):
        setattr(c, f, '')
    return doc


AN='/home/user/-Diagnostic-performance-of-intestinal-malrotation/analysis/'
OUT='/home/user/-Diagnostic-performance-of-intestinal-malrotation/'
D=json.load(open(AN+'tables123.json')); D.update(json.load(open(AN+'tables456.json')))
D.update(json.load(open(AN+'table3_final.json')))
D.update(json.load(open(AN+'tables24_final.json')))
T={'1':D['T1'],'2':D['T2'],'3':D['T3'],'4':D['T4']}
TITLES={
 '1':('Table 1','Characteristics of the 465 children with surgically confirmed intestinal malrotation, overall and according to which preoperative index test they received',
   'Groups overlap, because a child could receive more than one index test; columns therefore do not sum to the cohort total. Age at operation was calculated as age at admission for the operative encounter plus the interval from admission to operation. Presenting features were extracted from admission records by text search and are documentation rates, not verified prevalences. The rightmost column shows the 55 children who received none of the three index tests; all of them nonetheless had other preoperative imaging (Fig. 1). IQR interquartile range, UGI upper gastrointestinal.'),
 '2':('Table 2','Report-level detection of intestinal malrotation among surgically confirmed children, with the certainty of the wording used in positive conclusions',
   'Wilson 95% confidence intervals. Denominators differ between modalities and are drawn from overlapping but non-identical, indication-selected groups of children; the rates are not directly comparable between modalities and are not sensitivities. Certainty tiers were assigned from the conclusion text: definite (unqualified statement), probable ("most likely", "first consideration"), possible ("suspected", "cannot be excluded", "?"). The final column repeats the detection rate after reclassifying all possible-tier conclusions as negative. Percentages for certainty tiers are of positive reports; detection rates are of all index reports of that modality. Contrast-enhanced and unenhanced CT were performed for different indications and in children of different ages, so their comparison is confounded and is presented as an exploratory subgroup only.'),
 '3':('Table 3','Documented content of the 119 routine gastrointestinal ultrasound index examinations, and report-level detection conditional on that content',
   'The index unit is the examination episode closest to operation, pooling all reports of that modality issued that day. Content was coded from the findings and conclusion text by pre-specified patterns (Online Resource 1). This is an audit of what was recorded and is a lower bound on what was performed: an element assessed but not documented cannot be distinguished from one never assessed. Technique elements count as documented whether the finding was normal or abnormal; the whirlpool sign counts only where it is not negated, and agrees with the separately adjudicated whirlpool variable in 118 of 119 examinations. Enteric fluid administration counts only examinations stating that fluid was given (oral contrast or nasogastric instillation); observed luminal fluid without a stated route was not counted. Booking categories are not mutually exclusive: 28 sessions were booked as both a gastrointestinal and a great-vessel study. Denominators of fewer than 10 are shown without a percentage.'),
 '4':('Table 4','Change in report-level detection between eras, and the examination-content variables associated with it',
   'Modality-specific logistic models. For ultrasound the content variable is whether the examination session included an abdominal great-vessel study; this is an ordering and reporting label, not a record of technique. For CT the variable is intravenous contrast enhancement. Odds ratios and average marginal effects are both shown because a conditional odds ratio attenuates when a predictive covariate is added even in the absence of mediation (non-collapsibility), so the marginal effect is the more interpretable measure of how much of the era difference the content variable accounts for. The content variables were not randomised and are themselves confounded by indication: a child suspected of volvulus is both more likely to be booked for a great-vessel study and more likely to have a whirlpool to find. These models are explanatory, not causal. Sensitivity analyses moving the era boundary from 2019 to 2022 are in Online Resource 2.'),
}
FIGS={'1':(OUT+'Fig1_study_flow.png',6.4),'2':(OUT+'Fig2_detection_by_modality.png',6.4),'3':(OUT+'Fig3_ultrasound_report_audit.png',6.6)}
doc=docx.Document()
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(11)
for s in doc.sections: s.left_margin=s.right_margin=Inches(1.0)
def para(text,style=None,size=None,italic=False,space_after=8):
    p=doc.add_paragraph(style=style)
    for pt in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)',text):
        if not pt: continue
        if pt.startswith('**') and pt.endswith('**'): r=p.add_run(pt[2:-2]); r.bold=True
        elif pt.startswith('*') and pt.endswith('*') and len(pt)>2: r=p.add_run(pt[1:-1]); r.italic=True
        else: r=p.add_run(pt)
        if italic: r.italic=True
        if size: r.font.size=Pt(size)
    p.paragraph_format.space_after=Pt(space_after); return p
def add_table(k):
    tag,title,foot=TITLES[k]; data=T[k]
    p=para(f'{tag}. {title}'); p.runs[0].bold=True
    make_three_line_table(doc,data)
    para(foot,italic=True,size=8.5)
def add_fig(k):
    path,w=FIGS[k]
    doc.add_picture(path,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
text=''.join(open(f).read()+'\n' for f in ['p1.md','p2.md','p3.md'])
for ln in text.split('\n'):
    ln=ln.strip()
    if not ln: continue
    if ln.startswith('#T '): para(ln[3:],style='Title')
    elif ln.startswith('#H1 '): doc.add_heading(ln[4:],level=1)
    elif ln.startswith('#H2 '): doc.add_heading(ln[4:],level=2)
    elif ln.startswith('#N '): para(ln[3:])
    elif ln.startswith('#R '): para(ln[3:],size=10,space_after=3)
    elif ln.startswith('#TAB'): add_table(ln[4:])
    elif ln.startswith('#FIG'): add_fig(ln[4:])
    else: para(ln)
_blank_props(doc).save(OUT+'诊断效能_英文稿_InsightsIntoImaging投稿版_v4.docx'); print('saved')
