# -*- coding: utf-8 -*-
"""Shared three-line (open) table formatter for the manuscript and Online
Resource builders.

Journal convention for scientific tables (NEJM/JAMA/AMA style, and what Springer
journals including Insights into Imaging expect) is an "open" table: a rule above
the header, a rule below the header, a rule at the foot, and nothing else — no
vertical rules, no rules between data rows. This is the same convention as the
Chinese academic 三线表. python-docx has no built-in style for this, so every
border is set explicitly via OXML rather than relying on a Word table style
(built-in styles like "Table Grid" or "Light Grid Accent 1" draw a full grid and
cannot be trusted to render borderless even when overridden per cell).
"""
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

THICK = {'sz': '12', 'val': 'single', 'color': '000000'}   # ~1.5pt: top and bottom rules
THIN = {'sz': '6', 'val': 'single', 'color': '000000'}      # ~0.75pt: the rule under the header
NONE = {'sz': '0', 'val': 'nil', 'color': 'auto'}


def _set_cell_borders(cell, **edges):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
    for edge, spec in edges.items():
        etag = qn(f'w:{edge}')
        el = borders.find(etag)
        if el is None:
            el = OxmlElement(f'w:{edge}')
            borders.append(el)
        for k, v in spec.items():
            el.set(qn(f'w:{k}'), v)


def make_three_line_table(doc, data, header_rows=1, font_size=8.5,
                           font_name='Times New Roman'):
    """Build an open/three-line table. `data` is a list of row lists; the first
    `header_rows` rows are bold, with the header-underline rule beneath them."""
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.style = None  # no built-in style — every border is explicit, below
    tblPr = t._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        for k, v in NONE.items():
            el.set(qn(f'w:{k}'), v)
        tblBorders.append(el)
    tblPr.append(tblBorders)

    n_rows = len(data)
    for i, row in enumerate(data):
        for j, c in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ''
            r = cell.paragraphs[0].add_run(str(c))
            r.font.size = Pt(font_size)
            r.font.name = font_name
            if i < header_rows:
                r.bold = True
            edges = {'left': NONE, 'right': NONE, 'top': NONE, 'bottom': NONE}
            if i == 0:
                edges['top'] = THICK
            if i == header_rows - 1:
                edges['bottom'] = THIN
            if i == n_rows - 1:
                edges['bottom'] = THICK
            _set_cell_borders(cell, **edges)
            pf = cell.paragraphs[0].paragraph_format
            pf.space_after = Pt(2)
            pf.space_before = Pt(2)
    return t
