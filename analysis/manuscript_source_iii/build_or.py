import docx, json, re, os
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
AN='/tmp/claude-0/-home-user--Diagnostic-performance-of-intestinal-malrotation/4998e6b2-e14c-57b0-80b5-41a1fb987d43/scratchpad/an/'
MS='/tmp/claude-0/-home-user--Diagnostic-performance-of-intestinal-malrotation/4998e6b2-e14c-57b0-80b5-41a1fb987d43/scratchpad/ms/'
OUT='/home/user/-Diagnostic-performance-of-intestinal-malrotation/'
D=json.load(open(AN+'tables123.json')); D.update(json.load(open(AN+'tables456.json')))
D.update(json.load(open(AN+'or_tables.json'))); D.update(json.load(open(AN+'or_h.json')))
D.update(json.load(open(AN+'or3_pooled.json'))); D.update(json.load(open(AN+'or_sens.json')))
TAB={'H':('Table S1','Report-content audit patterns',D['H']),
     'S1':('Table S2','Distribution of algorithmic labels and certainty tiers, by modality',D['S1']),
     'G':('Table S3','Between-modality comparison of report-level detection (generalised estimating equation)',D['T4']),
     'P':('Table S4','Report-level detection in the selected subgroup receiving all three examinations, overall and restricted to examinations performed close together in time',D['T5']),
     'S2':('Table S5','Report-level detection stratified by midgut volvulus and by age category',D['S2']),
     'S2B':('Table S6','Detection of a volvulus-specific sign among children with surgically confirmed midgut volvulus',D['S2b']),
     'S3':('Table S7','Documented content of the CT and upper gastrointestinal series index examinations',D['S3']),
     'S4':('Table S8','Ultrasound temporal model with the era boundary placed at 2019, 2020, 2021 and 2022',D['S4']),
     'S5':('Table S9','Ultrasound content audit taking the earliest rather than the closest preoperative examination episode as the index unit',D['S5']),
     'S6':('Table S10','Prevalence of midgut volvulus under the primary and the restricted definition',D['S6'])}
def make(src,outfile,figs=None):
    doc=docx.Document()
    st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(11)
    for s in doc.sections: s.left_margin=s.right_margin=Inches(1.0)
    def para(t,style=None,size=None,italic=False):
        p=doc.add_paragraph(style=style)
        for pt in re.split(r'(\*\*[^*]+\*\*)',t):
            if not pt: continue
            if pt.startswith('**') and pt.endswith('**'): r=p.add_run(pt[2:-2]); r.bold=True
            else: r=p.add_run(pt)
            if italic: r.italic=True
            if size: r.font.size=Pt(size)
        p.paragraph_format.space_after=Pt(8); return p
    def table(key):
        tag,title,data=TAB[key]
        p=para(f'{tag}. {title}'); p.runs[0].bold=True
        t=doc.add_table(rows=len(data),cols=len(data[0])); t.style='Light Grid Accent 1'
        for i,row in enumerate(data):
            for j,c in enumerate(row):
                cell=t.cell(i,j); cell.text=''
                r=cell.paragraphs[0].add_run(str(c)); r.font.size=Pt(8.5); r.font.name='Times New Roman'
                if i==0: r.bold=True
        para('')
    for ln in open(src).read().split('\n'):
        ln=ln.strip()
        if not ln: continue
        if ln.startswith('#T '): para(ln[3:],style='Title')
        elif ln.startswith('#H1 '): doc.add_heading(ln[4:],level=1)
        elif ln.startswith('#N '): para(ln[3:])
        elif ln.startswith('#FIG'):
            doc.add_picture(OUT+'FigS1_paired_subgroup.png',width=Inches(6.2)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif ln.startswith('#TAB'): table(ln[4:])
        else: para(ln)
    doc.save(OUT+outfile); print('saved',outfile)
make(MS+'or1.md','Online_Resource_1_NLP_and_report_audit.docx')
make('or2_iii.md','Online_Resource_2_models_and_subgroups.docx')
make('or3_iii.md','Online_Resource_3_CT_and_UGI_content_audit.docx')
