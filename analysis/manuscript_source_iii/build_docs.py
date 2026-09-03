import docx, re, sys, io, os
from docx.shared import Pt, Inches
OUT='/home/user/-Diagnostic-performance-of-intestinal-malrotation/'
JOBS=[('cover.md','CoverLetter_InsightsIntoImaging.docx','Times New Roman',11),
      ('opsheet.md','投稿操作单_InsightsIntoImaging.docx','DengXian',10.5),
      ('cn2.md','投稿说明_InsightsIntoImaging_中文.docx','DengXian',10.5)]
for src,dst,font,size in JOBS:
    doc=docx.Document()
    st=doc.styles['Normal']; st.font.name=font; st.font.size=Pt(size)
    for s in doc.sections: s.left_margin=s.right_margin=Inches(1.0)
    def para(text,style=None,sz=None,italic=False,space_after=8):
        p=doc.add_paragraph(style=style)
        for pt in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)',text):
            if not pt: continue
            if pt.startswith('**') and pt.endswith('**'): r=p.add_run(pt[2:-2]); r.bold=True
            elif pt.startswith('*') and pt.endswith('*') and len(pt)>2: r=p.add_run(pt[1:-1]); r.italic=True
            else: r=p.add_run(pt)
            if italic: r.italic=True
            if sz: r.font.size=Pt(sz)
        p.paragraph_format.space_after=Pt(space_after); return p
    for ln in io.open(os.path.join(os.path.dirname(os.path.abspath(__file__)),src),encoding='utf-8').read().split('\n'):
        ln=ln.strip()
        if not ln: continue
        if ln.startswith('#T '): para(ln[3:],style='Title')
        elif ln.startswith('#H1 '): doc.add_heading(ln[4:],level=1)
        elif ln.startswith('#H2 '): doc.add_heading(ln[4:],level=2)
        elif ln.startswith('#N '): para(ln[3:])
        elif ln.startswith('#R '): para(ln[3:],sz=size-1,space_after=3)
        else: para(ln)
    doc.save(OUT+dst); print('saved',dst)
