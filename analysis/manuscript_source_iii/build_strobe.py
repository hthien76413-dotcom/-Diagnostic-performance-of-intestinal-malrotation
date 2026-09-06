# -*- coding: utf-8 -*-
"""Build STROBE_checklist.docx from strobe.json, with an open (three-line) table."""
import docx, json, os
from docx.shared import Pt, Inches
from threeline import make_three_line_table

D = os.path.dirname(os.path.abspath(__file__))
OUT = '/home/user/-Diagnostic-performance-of-intestinal-malrotation/'


def _blank_props(doc):
    c = doc.core_properties
    for f in ('author', 'last_modified_by', 'title', 'subject', 'comments',
              'category', 'keywords', 'content_status', 'identifier',
              'language', 'version'):
        setattr(c, f, '')
    return doc


rows = json.load(open(os.path.join(D, 'strobe.json')))['STROBE']

doc = docx.Document()
st = doc.styles['Normal']; st.font.name = 'Times New Roman'; st.font.size = Pt(10)
for s in doc.sections: s.left_margin = s.right_margin = Inches(0.7)


def para(text, bold=None, italic=None, size=None, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if size: r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


para("STROBE Statement — checklist of items that should be included in reports of cohort studies",
     bold=True, size=13)
para('Manuscript: "Routine ultrasound reports for intestinal malrotation rarely document duodenal '
     'landmarks: an audit of 740 preoperative index examinations in 410 surgically confirmed children"',
     size=10)
para("The study is a retrospective, single-centre, case-only report audit and is reported according "
     "to STROBE for cohort studies. STARD is deliberately not claimed: the cohort contains no "
     "test-negative children, so the design does not meet the definition of a diagnostic accuracy "
     "study and specificity, predictive values and any ranking of the three modalities are not "
     "estimable. Locations are given by manuscript section rather than by page, since pagination "
     "changes with the journal's template.", italic=True, size=9.5)

make_three_line_table(doc, rows, font_size=8.5)
doc.add_paragraph()

para("Note: An Explanation and Elaboration article discusses each checklist item and gives "
     "methodological background and published examples of transparent reporting. The STROBE "
     "checklist is best used in conjunction with that article, freely available on the STROBE "
     "website (https://www.strobe-statement.org).", italic=True, size=8.5, space_after=0)

_blank_props(doc)
doc.save(OUT + 'STROBE_checklist.docx')
print('saved STROBE_checklist.docx')
